import os
import json as json_module
from typing import Tuple, Dict, Any
from datetime import datetime

from pypdf import PdfReader
from docx import Document
import openpyxl

from app.services.ocr import OCRService


class DocumentProcessor:
    """Process various document types and extract text content"""

    def __init__(self, enable_ocr: bool = True):
        self.enable_ocr = enable_ocr
        self.ocr_service = OCRService() if enable_ocr else None

    def process(self, file_path: str, original_filename: str) -> Tuple[str, Dict[str, Any]]:
        """
        Process a document and extract text content with metadata

        Returns:
            Tuple of (text_content, metadata)
        """
        file_ext = os.path.splitext(file_path)[1].lower()
        file_size = os.path.getsize(file_path)

        metadata = {
            "filename": original_filename,
            "file_type": file_ext[1:],
            "file_size": file_size,
            "processed_at": datetime.utcnow().isoformat()
        }

        if file_ext == ".pdf":
            text, extra_metadata = self._process_pdf(file_path)
        elif file_ext == ".docx":
            text, extra_metadata = self._process_docx(file_path)
        elif file_ext == ".xlsx":
            text, extra_metadata = self._process_xlsx(file_path)
        elif file_ext == ".txt":
            text, extra_metadata = self._process_txt(file_path)
        elif file_ext == ".json":
            text, extra_metadata = self._process_json(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")

        metadata.update(extra_metadata)
        return text, metadata

    def _process_pdf(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Process PDF file"""
        reader = PdfReader(file_path)
        text_parts = []
        page_texts = []
        ocr_used = False

        for page_num, page in enumerate(reader.pages, 1):
            page_text = page.extract_text() or ""

            # If no text extracted and OCR is enabled, try OCR
            if not page_text.strip() and self.ocr_service and self.enable_ocr:
                try:
                    page_text = self.ocr_service.ocr_pdf_page(file_path, page_num - 1)
                    ocr_used = True
                except Exception:
                    pass

            if page_text.strip():
                page_texts.append({
                    "page": page_num,
                    "text": page_text
                })
                text_parts.append(f"[Page {page_num}]\n{page_text}")

        metadata = {
            "num_pages": len(reader.pages),
            "ocr_used": ocr_used,
            "page_data": page_texts
        }

        # Extract PDF metadata
        if reader.metadata:
            if reader.metadata.title:
                metadata["title"] = reader.metadata.title
            if reader.metadata.author:
                metadata["author"] = reader.metadata.author
            if reader.metadata.creation_date:
                metadata["creation_date"] = str(reader.metadata.creation_date)

        return "\n\n".join(text_parts), metadata

    def _process_docx(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Process Word document"""
        doc = Document(file_path)
        text_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

        metadata = {
            "num_paragraphs": len(doc.paragraphs),
            "num_tables": len(doc.tables)
        }

        # Extract document properties
        if doc.core_properties:
            if doc.core_properties.title:
                metadata["title"] = doc.core_properties.title
            if doc.core_properties.author:
                metadata["author"] = doc.core_properties.author

        return "\n\n".join(text_parts), metadata

    def _process_xlsx(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Process Excel file"""
        workbook = openpyxl.load_workbook(file_path, data_only=True)
        text_parts = []
        sheet_names = []

        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            sheet_names.append(sheet_name)
            sheet_text = [f"[Sheet: {sheet_name}]"]

            for row in sheet.iter_rows():
                row_values = []
                for cell in row:
                    if cell.value is not None:
                        row_values.append(str(cell.value))
                if row_values:
                    sheet_text.append(" | ".join(row_values))

            if len(sheet_text) > 1:
                text_parts.append("\n".join(sheet_text))

        metadata = {
            "num_sheets": len(workbook.sheetnames),
            "sheet_names": sheet_names
        }

        return "\n\n".join(text_parts), metadata

    def _process_txt(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Process text file"""
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()

        metadata = {
            "num_lines": content.count("\n") + 1,
            "num_characters": len(content)
        }

        return content, metadata

    def _process_json(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Process JSON file as raw text (fallback for non-pre-chunked JSON)"""
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()

        data = None
        try:
            data = json_module.loads(content)
            text = json_module.dumps(data, ensure_ascii=False, indent=2)
        except json_module.JSONDecodeError:
            text = content

        metadata = {
            "num_characters": len(text),
        }

        if data is not None:
            metadata["json_type"] = type(data).__name__
            if isinstance(data, list):
                metadata["num_items"] = len(data)

        return text, metadata
